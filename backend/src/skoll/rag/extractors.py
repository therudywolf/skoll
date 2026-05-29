"""File -> text extractors.

Issue: phase-1.6.

Adapted from vendor/ForestOptiLM/file_extractors.py, narrowed to the formats the
Skoll RAG pipeline needs and the dependencies pinned in pyproject.toml:

  - plain text / source code / markup  -> decoded directly
  - .pdf                               -> pypdf
  - .docx                              -> python-docx
  - .xlsx                              -> openpyxl
  - .html / .htm                       -> trafilatura, falling back to BeautifulSoup
  - images                             -> skipped (the vision tool handles those)

Unknown / unsupported extensions raise ``skoll.errors.ToolExecutionError`` (the
Golden Rules forbid raising bare ``Exception``).
"""

from __future__ import annotations

import re
from pathlib import Path

from skoll.errors import ToolExecutionError

# Source / markup / config files that are read verbatim as UTF-ish text.
_TEXT_SUFFIXES: frozenset[str] = frozenset(
    {
        # Plain text + docs
        ".txt",
        ".md",
        ".markdown",
        ".rst",
        ".log",
        # Config / data markup
        ".ini",
        ".cfg",
        ".conf",
        ".toml",
        ".env",
        ".properties",
        ".xml",
        ".yaml",
        ".yml",
        ".json",
        ".csv",
        ".tsv",
        # Code
        ".py",
        ".js",
        ".ts",
        ".tsx",
        ".jsx",
        ".mjs",
        ".cjs",
        ".c",
        ".cpp",
        ".cc",
        ".h",
        ".hpp",
        ".java",
        ".kt",
        ".scala",
        ".go",
        ".rs",
        ".rb",
        ".php",
        ".cs",
        ".swift",
        ".sql",
        ".sh",
        ".bash",
        ".zsh",
        ".bat",
        ".cmd",
        ".ps1",
        ".r",
        ".lua",
        ".dart",
    }
)

# Images are handled by the vision tool, not the text RAG path.
_IMAGE_SUFFIXES: frozenset[str] = frozenset(
    {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp", ".tif", ".tiff", ".svg", ".ico"}
)

_HTML_SUFFIXES: frozenset[str] = frozenset({".html", ".htm", ".xhtml"})

# Decode candidates, ordered most→least likely. Mirrors ForestOptiLM.
_ENCODINGS = ("utf-8-sig", "utf-8", "cp1251", "cp1252", "latin-1")


def _decode(raw: bytes) -> str:
    """Best-effort decode of arbitrary file bytes to ``str``."""
    if raw[:2] in (b"\xff\xfe", b"\xfe\xff"):
        return raw.decode("utf-16", errors="replace")
    for enc in _ENCODINGS:
        try:
            return raw.decode(enc)
        except (UnicodeDecodeError, LookupError):
            continue
    return raw.decode("utf-8", errors="replace")


def _extract_text(path: Path) -> str:
    return _decode(path.read_bytes())


def _extract_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    parts: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            parts.append(text)
    return "\n".join(parts)


def _extract_docx(path: Path) -> str:
    from docx import Document

    document = Document(str(path))
    lines: list[str] = [p.text for p in document.paragraphs if p.text.strip()]
    # Tables carry data that paragraphs miss — flatten each row to a tab line.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                lines.append("\t".join(cells))
    return "\n".join(lines)


def _extract_xlsx(path: Path) -> str:
    from openpyxl import load_workbook

    # read_only + data_only: stream cells, return computed values not formulas.
    workbook = load_workbook(filename=str(path), read_only=True, data_only=True)
    try:
        sheets: list[str] = []
        for worksheet in workbook.worksheets:
            rows: list[str] = []
            for row in worksheet.iter_rows(values_only=True):
                cells = ["" if value is None else str(value) for value in row]
                if any(cell.strip() for cell in cells):
                    rows.append("\t".join(cells))
            if rows:
                sheets.append(f"# Sheet: {worksheet.title}\n" + "\n".join(rows))
        return "\n\n".join(sheets)
    finally:
        workbook.close()


def _extract_html(path: Path) -> str:
    raw = path.read_bytes()
    markup = _decode(raw)

    # Primary: trafilatura — boilerplate-stripped main content.
    import trafilatura

    # trafilatura is untyped (Any); it returns the extracted text or None.
    extracted: str | None = trafilatura.extract(markup)
    if extracted and extracted.strip():
        return str(extracted)

    # Fallback: BeautifulSoup full-text (trafilatura returns None on fragments).
    try:
        from bs4 import BeautifulSoup

        text: str = BeautifulSoup(markup, "html.parser").get_text(separator="\n", strip=True)
        return text
    except Exception:  # pragma: no cover - bs4 is a hard dependency; defensive only
        return re.sub(r"<[^>]+>", " ", markup)


def extract(path: Path) -> str:
    """Return a plain-text representation of ``path``.

    Raises:
        ToolExecutionError: the extension is an image (vision-only), is otherwise
            unsupported, or the underlying parser failed.
    """
    path = Path(path)
    if not path.is_file():
        raise ToolExecutionError(f"Not a file: {path}")

    suffix = path.suffix.lower()

    if suffix in _IMAGE_SUFFIXES:
        raise ToolExecutionError(
            f"Image file {path.name!r} is not extractable as text; use the vision tool."
        )

    try:
        if suffix in _HTML_SUFFIXES:
            return _extract_html(path)
        if suffix == ".pdf":
            return _extract_pdf(path)
        if suffix == ".docx":
            return _extract_docx(path)
        if suffix == ".xlsx":
            return _extract_xlsx(path)
        if suffix in _TEXT_SUFFIXES:
            return _extract_text(path)
    except ToolExecutionError:
        raise
    except Exception as exc:  # parser-specific failures are normalised to ToolExecutionError
        raise ToolExecutionError(f"Failed to extract {path.name!r}: {exc}") from exc

    raise ToolExecutionError(f"Unsupported file type for RAG extraction: {suffix or path.name!r}")

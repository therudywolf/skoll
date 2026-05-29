"""Round-trip tests for file -> text extractors (Issue 1.6).

Static fixtures (.txt/.md/.html) live in tests/fixtures/rag/. Binary office
formats (.docx/.xlsx) are authored on the fly so the repo carries no opaque blobs.
PDF content extraction needs a real PDF, so only the dispatch/error path is tested
here; a content round-trip is marked ``integration``.
"""

from __future__ import annotations

from pathlib import Path

import pytest
from skoll.errors import ToolExecutionError
from skoll.rag.extractors import extract


@pytest.fixture
def rag_fixtures(fixtures_dir: Path) -> Path:
    return fixtures_dir / "rag"


# --------------------------------------------------------------------------- #
# Plain text + markup
# --------------------------------------------------------------------------- #


def test_extract_txt(rag_fixtures: Path) -> None:
    text = extract(rag_fixtures / "sample.txt")
    assert "local-first agentic web IDE" in text
    assert "RAG pipeline indexes a workspace" in text


def test_extract_md(rag_fixtures: Path) -> None:
    text = extract(rag_fixtures / "sample.md")
    # Markdown is read verbatim — heading markers are preserved.
    assert "# Sample Document" in text
    assert "**markdown**" in text


def test_extract_html_strips_boilerplate(rag_fixtures: Path) -> None:
    text = extract(rag_fixtures / "sample.html")
    assert "quick brown fox" in text
    assert "real article content" in text
    # Tags are gone.
    assert "<p>" not in text and "<article>" not in text


def test_extract_unknown_text_code_suffix(tmp_path: Path) -> None:
    src = tmp_path / "snippet.py"
    src.write_text("def hello() -> str:\n    return 'hi'\n", encoding="utf-8")
    text = extract(src)
    assert "def hello()" in text


# --------------------------------------------------------------------------- #
# DOCX (authored in-test)
# --------------------------------------------------------------------------- #


def test_extract_docx(tmp_path: Path) -> None:
    from docx import Document

    doc = Document()
    doc.add_paragraph("First paragraph of the docx fixture.")
    doc.add_paragraph("Second paragraph with more words.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "alpha"
    table.cell(1, 1).text = "42"
    path = tmp_path / "doc.docx"
    doc.save(str(path))

    text = extract(path)
    assert "First paragraph of the docx fixture." in text
    assert "Second paragraph with more words." in text
    # Table cells are flattened to tab-separated lines.
    assert "Name\tValue" in text
    assert "alpha\t42" in text


# --------------------------------------------------------------------------- #
# XLSX (authored in-test)
# --------------------------------------------------------------------------- #


def test_extract_xlsx(tmp_path: Path) -> None:
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Data"
    ws.append(["city", "population"])
    ws.append(["Reykjavik", 131000])
    ws.append(["Oslo", 700000])
    path = tmp_path / "book.xlsx"
    wb.save(str(path))

    text = extract(path)
    assert "# Sheet: Data" in text
    assert "city\tpopulation" in text
    assert "Reykjavik\t131000" in text
    assert "Oslo\t700000" in text


# --------------------------------------------------------------------------- #
# Dispatch / error paths
# --------------------------------------------------------------------------- #


def test_image_suffix_is_rejected(tmp_path: Path) -> None:
    img = tmp_path / "pic.png"
    img.write_bytes(b"\x89PNG\r\n\x1a\n")  # PNG magic; content irrelevant
    with pytest.raises(ToolExecutionError, match="vision"):
        extract(img)


def test_unsupported_suffix_is_rejected(tmp_path: Path) -> None:
    blob = tmp_path / "thing.unknownext"
    blob.write_bytes(b"\x00\x01\x02")
    with pytest.raises(ToolExecutionError, match="Unsupported file type"):
        extract(blob)


def test_missing_file_is_rejected(tmp_path: Path) -> None:
    with pytest.raises(ToolExecutionError, match="Not a file"):
        extract(tmp_path / "does_not_exist.txt")


def test_pdf_dispatch_normalises_parse_failure(tmp_path: Path) -> None:
    # We do not author a valid PDF (impractical without extra deps); a malformed
    # one proves the .pdf branch dispatches to pypdf and that parser errors are
    # normalised to ToolExecutionError rather than leaking a raw pypdf exception.
    bad = tmp_path / "broken.pdf"
    bad.write_bytes(b"%PDF-1.4\nnot a real pdf body\n")
    with pytest.raises(ToolExecutionError, match=r"broken\.pdf"):
        extract(bad)


@pytest.mark.integration
def test_extract_pdf_content() -> None:
    """Round-trip a real PDF. Skipped in CI: requires a committed/binary PDF
    fixture, which we avoid authoring without a PDF-writing dependency."""
    pytest.skip("PDF content round-trip needs a real PDF fixture / writer dependency")
